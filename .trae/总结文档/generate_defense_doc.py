import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.4

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)

def add_body(text):
    doc.add_paragraph(text)

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = 'Consolas'

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

def add_note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 50, 50)

# ==================== 正文 ====================

add_title('智能宿舍管理系统\n答辩准备手册')

# ==================== 重要提醒 ====================
add_heading('⚠️ 论文与实际代码的关键差异（答辩前务必核对）', level=2)
add_note('以下是论文中与你实际代码不一致的地方，老师可能发现，你必须清楚：')

add_table(['序号', '论文写的', '实际代码是', '文件位置'], [
    ['1', '前端用 Vuex 管理状态', '用 Pinia（Vue 3 官方推荐）', 'dormitory-web/src/store/user.js 第1行：import { defineStore } from "pinia"'],
    ['2', 'Token 存在 localStorage', 'Token 存在 sessionStorage（关浏览器即清除）', 'dormitory-web/src/store/user.js，如 sessionStorage.getItem("admin_token")'],
])

add_body('答辩时如果老师指出这些差异：大方承认"论文里写的 Vuex/localStorage 是早期方案的描述，实际开发中改用了 Pinia/sessionStorage，因为 Pinia 是 Vue 3 官方推荐的状态管理库，sessionStorage 比 localStorage 更安全，关闭浏览器就自动退出。"')

add_heading('📌 论文各章节对应代码位置速查', level=2)

add_table(['论文章节', '内容', '对应的实际代码文件'], [
    ['3.3 用例图分析', '三类角色用例', '功能文档：dormitory_system.sql（29张表DDL）'],
    ['4.1 系统架构', 'B/S + MVC三层', 'dormitory-admin/src/main/java/com/dormitory/DormitoryApplication.java（启动类）\ndormitory-admin/src/main/resources/application.yml（配置）'],
    ['4.3 数据库逻辑结构', 'E-R图', 'dormitory_system.sql（项目根目录，完整建表语句）'],
    ['4.4 系统数据表设计', '29张表字段说明', 'dormitory-admin/src/main/java/com/dormitory/entity/（每个表对应一个Entity类）'],
    ['5.1.1 登录认证', 'BCrypt + JWT', 'AuthServiceImpl.java + JwtUtil.java + JwtAuthenticationFilter.java'],
    ['5.1.2 在线报修', '状态机流程', 'RepairController.java + RepairCommentController.java'],
    ['5.1.3 换寝申请', '多表联动更新', 'RoomChangeController.java'],
    ['5.1.4 智能客服', 'DeepSeek AI', 'ChatController.java'],
    ['5.2.1 打卡管理', '晚归打卡+补卡', 'CheckInController.java'],
    ['5.2.2 水电费管理', '阈值预警', 'UtilityController.java + UtilityThresholdController.java'],
    ['5.2.3 文明宿舍', '自动评分排名', 'CivilizedDormitoryController.java'],
    ['5.3.1 用户管理', '三类用户CRUD', 'StudentController.java + ManagerController.java + AdminController.java'],
    ['5.3.2 数据统计与导入导出', 'POI Excel', 'StatisticsController.java'],
    ['6 系统测试', '功能/性能/安全/兼容性', '.trae/specs/安全性测试规范/ + .trae/specs/性能测试规范/'],
])

# ==================== 第1部分：论文相关提问 ====================
add_heading('第一部分：基于论文内容的常见提问', level=1)

# Q1
add_heading('Q1：为什么选这个课题？有什么实际意义？', level=2)
add_body('高校宿舍管理长期靠人工——报修要跑腿、水电费要上门抄表、查信息要去管理处。这套系统把宿舍管理全流程线上化，减少人工干预。论文第1章写了：目前高校招生规模扩大，传统方式"信息不及时，数据统计困难，报修效率低"。系统让学生在线办业务、宿管高效管楼栋、管理员有数据做决策。')

# Q2
add_heading('Q2：为什么选 Spring Boot + Vue 这套技术栈？', level=2)
add_body('论文第2章有详细说明。Spring Boot 省去了传统 Spring 的繁琐配置，启动器和自动配置能快速集成 MyBatis-Plus、Spring Security。Vue 3 组件化开发效率高，Element Plus 组件丰富省得自己写样式。前后端分离，各自独立部署和维护。MySQL 免费开源，对学生项目够用。')

# Q3
add_heading('Q3：你的系统架构是怎样的？', level=2)
add_body('论文第4章、图4-1。B/S架构，前后端分离。前端 Vue 3 + Vite（5173端口），后端 Spring Boot（8080端口），通过 HTTP RESTful API 交互，JSON 格式。后端内部是三层：Controller 接收请求 → Service 处理业务 → Mapper 操作数据库。用 Spring Security + JWT 做身份认证和权限控制。')
add_body('如果老师问"B/S和C/S有什么区别"：B/S 用浏览器访问不用装软件，C/S 要装客户端（如微信、QQ）。本系统选 B/S 是因为学生宿管用浏览器打开就能用。')

# Q4
add_heading('Q4：数据库有多少张表？怎么设计的？', level=2)
add_body('论文第4.4节列出了完整的表设计。共29张表，分三大类：')
add_body('① 用户类：admin、dormitory_manager、student —— 三种角色分别存')
add_body('② 宿舍资源类：building（楼）→ room（房间）→ bed（床位），三级关联。building_id 关联 building 表，room_id 关联 room 表')
add_body('③ 业务类：repair（报修）、utility_bill（水电费）、check_in（打卡）、room_change（换寝）、visitor（访客）、emergency_help（紧急求助）、lost_and_found（失物招领）、health_check（卫生检查）、chat_session/chat_message（客服）、notice（公告）、message（消息通知）、civilized_dormitory（文明宿舍）')
add_body('④ 辅助类：operation_log、login_log（日志）')
add_body('如果老师追问表之间的关系：宿舍楼 1:N 宿舍，宿舍 1:N 床位，学生入住一个床位。报修单关联学生和宿舍。')

# Q5
add_heading('Q5：论文写了哪些测试？结果怎么样？', level=2)
add_body('论文第6章写了四种测试：')
add_body('① 功能测试：逐个功能走了一遍，学生登录、报修、缴费等核心流程通过了')
add_body('② 性能测试：做了并发测试，系统能支持多人同时在线')
add_body('③ 安全性测试：SQL注入测试、XSS测试、JWT Token篡改测试、越权访问测试。未登录会返回401/403，SQL注入和XSS攻击无效')
add_body('④ 兼容性测试：Chrome、Edge 浏览器正常，Windows 10/11 系统正常')
add_body('测试文档目录：.trae/specs/安全性测试规范/、.trae/specs/性能测试规范/')

# Q6
add_heading('Q6：论文中提到的"技术难点"有哪些？你怎么解决的？', level=2)
add_body('论文各功能实现章节都写了技术难点，主要三个：')
add_body('① Token有效期管理：设了7天过期，兼顾安全性和体验。前端在 request.js 用响应拦截器捕获401自动跳登录页。')
add_body('② 报修单状态流转：5个状态（待处理→已接单→维修中→已完成/已取消），只能按顺序流转，不能跳状态。Constant.java 定义了状态常量。')
add_body('③ 换寝多表联动：审批通过后要同步更新 bed（释放+占用）、room（人数±1）、student（更新住宿信息）三张表。在 RoomChangeController.java 里一个方法完成所有更新。')

# Q7
add_heading('Q7：论文的创新点是什么？', level=2)
add_body('论文里强调了几个：')
add_body('① DeepSeek AI智能客服：7x24小时自动回答学生问题（ChatController.java）')
add_body('② 水电费先用后付+超限预警：按宿舍类型（4/6人间）区别定价，超限自动推送预警通知（UtilityController.java）')
add_body('③ Excel批量导入导出：学生信息、水电费、打卡数据支持批量操作（StatisticsController.java + Apache POI）')
add_body('④ 文明宿舍自动评分排名：根据卫生检查打分自动计算均分并排名（CivilizedDormitoryController.java）')

# ==================== 第2部分：技术细节快速定位 ====================
add_heading('第二部分：技术细节 & 代码快速定位', level=1)
add_body('以下按"老师可能怎么问 → 你怎么答 → 代码在哪"来组织。不需要解释代码逻辑，只需要知道在哪个文件、这个文件干什么。')

# T1
add_heading('T1：密码加密怎么做的？用什么算法？为什么选它？', level=2)
add_body('用的是 BCrypt 加密算法，Spring Security 自带的 BCryptPasswordEncoder。')
add_body('为什么选 BCrypt 而不是 MD5 或 SHA256：')
add_body('① MD5/SHA256 是"快哈希"，几微秒就能算完，黑客用 GPU 一秒能试几十亿次，暴力破解太快')
add_body('② BCrypt 是"慢哈希"，专门为密码存储设计的，默认计算 1024 轮，每次加密要 100-200 毫秒。对正常用户登录几乎无感，对黑客暴力破解是天堑')
add_body('③ BCrypt 每次加密自动生成随机"盐值"混入密码，两个用户设相同密码，数据库里密文也不同。黑客没法用彩虹表批量比对')
add_body('简单理解：MD5 像一把普通锁，钥匙对了秒开。BCrypt 像一把需要 100 道工序才能转动的锁——主人开锁等 0.1 秒无所谓，小偷要试 100 万个密码就得等到天荒地老。')
add_body('代码位置：SecurityConfig.java 定义 BCryptPasswordEncoder Bean → 新增用户时调 .encode("密码") 加密存库 → 登录时调 .matches(输入密码, 库中密文) 比对（不解密，只比对）。')
add_table(['文件', '路径', '作用'], [
    ['SecurityConfig.java', 'dormitory-admin/.../config/', 'BCryptPasswordEncoder Bean 定义，cost=10（1024轮）'],
    ['AuthServiceImpl.java', 'dormitory-admin/.../service/impl/', '登录验证：.matches(明文, 密文)'],
    ['StudentController.java\nManagerController.java\nAdminController.java', 'dormitory-admin/.../controller/', '新增用户时：.encode("123456")'],
])

# T2
add_heading('T2：JWT Token 认证流程是怎样的？', level=2)
add_body('三个文件配合：JwtUtil.java 生成和解析Token → JwtAuthenticationFilter.java 拦截每个请求验证Token → AuthServiceImpl.java 登录时调用生成。前端在 request.js 里统一带 Token（放 Authorization 头里），用 sessionStorage 存储，关浏览器即失效。')
add_table(['文件', '路径', '作用'], [
    ['JwtUtil.java', 'dormitory-admin/.../utils/', 'generateToken() 生成，parseToken() 解析'],
    ['JwtAuthenticationFilter.java', 'dormitory-admin/.../config/', '拦截请求，验证Token有效性'],
    ['AuthServiceImpl.java', 'dormitory-admin/.../service/impl/', '登录成功调 JwtUtil 生成 Token'],
    ['application.yml', 'dormitory-admin/src/main/resources/', 'jwt.expiration: 604800000（7天）'],
    ['request.js', 'dormitory-web/src/utils/', '前端统一带Token，401自动跳登录页'],
])

# T3
add_heading('T3：报修功能实现在哪？状态怎么流转的？', level=2)
add_body('状态在 Constant.java 里定义：0=待处理，1=已接单，2=维修中，3=已完成，4=已取消。')
add_body('学生提交报修（RepairController.java 的 submit 接口）→ 宿管接单（同一 Controller 的 accept 接口）→ 维修完标记完成 → 通知学生评价（RepairCommentController.java）。')
add_table(['文件', '作用'], [
    ['RepairController.java', '报修核心：提交、接单、完成、取消'],
    ['RepairCommentController.java', '学生评价（满意/一般/不满意）'],
    ['Constant.java', '状态常量定义'],
    ['student/repair/index.vue', '学生端报修页面'],
    ['manager/repair/index.vue', '宿管端报修处理页面'],
])

# T4
add_heading('T4：换寝功能怎么保证数据一致？', level=2)
add_body('RoomChangeController.java 里 approve 方法，审批通过后在一个事务里完成6步：①释放原床位(bed表状态→空闲) ②原宿舍人数-1 ③占用新床位 ④新宿舍人数+1 ⑤更新student表的roomId和bedNumber ⑥申请状态→已通过。所有操作在一个方法里，数据库事务保证原子性，要么全成功要么全回滚。')

# T5
add_heading('T5：AI 智能客服怎么接的？调了什么 API？', level=2)
add_body('ChatController.java 里 getAIResponse() 方法，用 Java 11 的 HttpClient 发送 POST 请求到 https://api.deepseek.com/v1/chat/completions，带上 API Key（配在 application.yml 的 deepseek.api-key）。会设 30 秒超时，失败了返回"AI 服务暂时不可用"。')
add_body('数据库用两张表：chat_session（会话）+ chat_message（消息），支持多轮对话。chat_type=1 是 AI客服，chat_type=2 是人工客服。')

# T6
add_heading('T6：水电费怎么算的？超限预警怎么做？', level=2)
add_body('UtilityThresholdController.java 里管理员配置不同宿舍类型（4/6人间）的单价和限额。UtilityController.java 里宿管录入水电用量时，系统自动读对应单价计算费用。checkAndCreateWarning() 方法检查是否超限，超了就创建预警记录并给该宿舍所有学生推送消息通知。学生端模拟支付：选微信/支付宝 → 弹二维码 → 15秒后自动标记已支付。')

# T7
add_heading('T7：权限控制怎么做的？学生能看到宿管的数据吗？', level=2)
add_body('数据层面的细粒度控制。JwtAuthenticationFilter 验证 Token 后把 userId 和 userType 放入上下文。Controller 里根据 userType 过滤数据：比如宿管调 GET /api/repair/list，只返回本楼栋的报修单（LambdaQueryWrapper 加 buildingId 条件）。学生只能看自己的，宿管看本楼栋，管理员看全部。')

# T8
add_heading('T8：Excel 导入导出用的什么库？怎么做的？', level=2)
add_body('Apache POI（poi-ooxml 5.2.5），依赖在 pom.xml。导入：前端 el-upload 传文件 → 后端接 MultipartFile → XSSFWorkbook 读 → 逐行解析 → 批量插入。导出：XSSFWorkbook 创建 → 写表头 → 循环写数据 → 设 Content-Disposition 响应头 → 浏览器自动下载。StatisticsController.java 和部分业务 Controller 里有对应方法。')

# T9 — N+1优化
add_heading('T9：学生列表查询做过什么性能优化？', level=2)
add_body('这是项目中最有价值的一个技术优化，解决了经典的 N+1 查询问题。')
add_body('优化前的问题：学生列表接口查询一次分页学生（1次SQL），然后循环为每个学生查床位，再为每个床位查房间、查楼栋。总共 1501 次数据库查询，页面加载 2 秒。')
add_body('优化方案：在 StudentController.java 的 list() 方法（第 188-236 行），改为批量查询 + 内存 Map 关联——')
add_body('① 只查当前页学生的床位（用 IN 条件，1次SQL）')
add_body('② 批量查所有相关房间（selectBatchIds，1次SQL）')
add_body('③ 批量查所有相关楼栋（selectBatchIds，1次SQL）')
add_body('④ 在内存里用 Map<Long, Room> 和 Map<Long, Building> 做关联组装')
add_body('效果：总查询从 1501 次降到 4 次，页面加载从 2 秒降到约 50ms。')
add_body('优化文档：项目根目录/学生列表查询性能优化方案.md')
add_body('核心代码：dormitory-admin/src/main/java/com/dormitory/controller/StudentController.java')
add_body('如果老师追问"整个系统都优化了吗"：核心 list 方法已优化，部分辅助方法后续可以继续优化。比如同一个 Controller 里第 479 行附近还有全表扫描床位的方法，是下一步优化的目标。')

# T10
add_heading('T10：application.yml 里配了什么？', level=2)
add_body('数据库连接（MySQL 3307，dormitory_system 库，HikariCP 最大50连接）、JWT 密钥和 7 天过期时间、DeepSeek API Key、文件上传最大 10MB、MyBatis 下划线转驼峰 + SQL 日志输出。文件位置：dormitory-admin/src/main/resources/application.yml。')

# T11
add_heading('T11：MyBatis-Plus 和普通 MyBatis 有什么区别？', level=2)
add_body('MyBatis 需要手写 SQL 和 XML 映射文件。MyBatis-Plus 增强了它：单表 CRUD 不用写 SQL（BaseMapper 自带 selectList、insert、updateById），LambdaQueryWrapper 拼条件避免字段名写错（如 lambdaQueryWrapper.eq(Student::getName, "张三")），内置分页插件（new Page<>(pageNum, pageSize)）。复杂联表查询才需要手写 SQL。')

# T12
add_heading('T12：Spring Security 在你项目里起了什么作用？', level=2)
add_body('SecurityConfig.java：① 放行 /api/auth/**（登录接口不拦截）② 其他接口必须认证。JwtAuthenticationFilter.java：每个请求到达后先解析 Token → 验证签名和有效期 → 把用户信息放入 SecurityContext → 放行。BCryptPasswordEncoder 的 Bean 也在这定义。')

# T13
add_heading('T13：跨域问题怎么解决的？', level=2)
add_body('后端 WebConfig.java 配了 CORS 跨域：允许前端 localhost:5173 的请求，允许带 Authorization 请求头。开发时 Vite 也配了 proxy 代理：/api 前缀的请求转给 localhost:8080。')

# ==================== 第3部分：演示准备 ====================
add_heading('第三部分：现场演示准备', level=1)

add_body('老师最可能让你演示的流程：')
add_body('① 登录：三个角色都登一遍，展示不同界面')
add_body('② 报修全流程：学生提交 → 宿管接单 → 维修完成 → 学生评价（一口气走完）')
add_body('③ 换寝：学生申请 → 宿管审批通过 → 学生信息自动更新')
add_body('④ AI 客服：学生问一个问题，展示 AI 回复')
add_body('⑤ 水电费：宿管录入 → 学生看到账单 → 扫码支付（模拟）')
add_body('⑥ Excel 导出：导出学生列表，下载打开看')

add_body('准备 CheckList：')
add_body('☐ 后端项目能正常启动（8080端口）')
add_body('☐ 前端项目能正常启动（5173端口）')
add_body('☐ MySQL 数据库正常运行（3307端口）')
add_body('☐ 三个角色的测试账号各备一个（学生/宿管/管理员）')
add_body('☐ DeepSeek API Key 有效，AI 客服能用')
add_body('☐ 提前打开好 TRAE，方便让老师看代码（你用的就是 TRAE 开发的，不用装别的编辑器）')

# ==================== 第4部分：答辩技巧 ====================
add_heading('第四部分：答辩应急技巧', level=1)

add_heading('如果不确定答案', level=2)
add_body('"老师，这部分我开发时主要关注了 A 和 B，您问的 C 我后续会继续研究完善。"——诚实但表明自己有思考。')

add_heading('如果被追问代码细节', level=2)
add_body('直接说文件名："这个功能在 XXController.java 里实现的，我可以打开给您看。"——知道代码在哪比背代码重要。')

add_heading('如果被质疑技术选型', level=2)
add_body('"选 Spring Boot 是因为它自动配置省时间，Vue 3 是因为组件化开发效率高。"——不要辩解，客观说理由。')

add_heading('如果老师指出论文错误', level=2)
add_body('"谢谢老师指出，论文里这部分写得不够准确，实际实现中是……"—大方承认，给出正确信息。特别是 Vuex vs Pinia、localStorage vs sessionStorage 这两个已知差异。')

add_heading('如果演示时系统挂了', level=2)
add_body('"可能是环境问题，我之前测试时一切正常。我可以先用代码给老师看这个功能的实现。"——提前准备好 IDE 里的代码，随时可以切过去展示。')

# 保存
output_path = r'C:\Users\32010\Desktop\AI项目\TRAE\.trae\答辩准备手册_v4.docx'
doc.save(output_path)
print(f'Done: {output_path}')
